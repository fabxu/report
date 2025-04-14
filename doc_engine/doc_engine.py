import re
from datetime import datetime
from typing import Any

from core import constant
from core.context.context import Context
from core.errors.error import Error, ErrorCode
from docx import Document

from core.resource.resource import ResType
from core.utils.log import logger
from doc_engine.image_render import ImageGen

PATTEN_ENV = "\#\w+[\w._]*\#"
PATTEN_PLUGIN = "\$\w+[\w._]*\$"
FLAG_PLUGIN = "$"


class DocEngin:
    def __init__(self):
        self._generators: dict = {}
        self._generators[ResType.IMAGE] = ImageGen()

    def createReport(self, templatePath: str, name: str, path: str) -> Error:
        if (name is not None):
            formatted_time = datetime.now().strftime(constant.FORMAT_TIME)
            self.reportPath = path + name + "_" + formatted_time + constant.SUFFIX_DOCX
            self.report = Document(templatePath)
            return Error(ErrorCode.SUCCESS, None)
        else:
            return Error(ErrorCode.NOT_SUPPORT, f"file name is not docx, name: {name}")

    def saveReport(self):
        self.report.save(self.reportPath)

    def _getPluginVar(self, content: str, startIndex: int) -> (int, str):
        start = content.index(FLAG_PLUGIN, startIndex)
        if start < 0:
            return constant.INVALID, None
        else:
            end = content.index(FLAG_PLUGIN, start + 1)
            if end < 0 and end - start > 1:
                return constant.INVALID, None
            else:
                return start, content[start: end + 1]

    def _processComplexContent(self, complexContent: dict, run, paragraph):
        start = 0
        content = run.text
        while (start >= 0) and len(content) > 0:
            start, var = self._getPluginVar(content, start)
            if start >= 0:
                if var in complexContent:
                    if start > 0:
                        self._copyRunProperty(run, paragraph, content[:start])
                    value = complexContent[var]
                    self._generators[value.getType()].genContent(self.report, value, run, paragraph)
                    content = content.replace(var, "", 1)
                    start = 0
                else:
                    start = start + len(var) - 1
        run.text = content

    def converValue2Str(self, value):
        if value is None:
            return repr(value)
        if not isinstance(value, str):
            return repr(value)
        return value

    def _processTable(self, ctx: Context):
        for i in range(0, len(self.report.tables)):
            table = self.report.tables[i]
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            content = run.text
                            envValues = self._parserEnv(content, PATTEN_ENV, ctx)
                            pluginValues = self._parserPlugin(content, PATTEN_PLUGIN, ctx)
                            if len(envValues) > 0 or len(pluginValues) > 0:
                                for key, value in envValues.items():
                                    value = self.converValue2Str(value)
                                    content = content.replace(key, value)
                                for key, value in pluginValues.items():
                                    if value.getType() == ResType.TEXT:
                                        content = content.replace(key, value.getData())
                                run.text = content

    def get_shading_color(self, run):
        """获取 Run 的字体背景色（高亮颜色）"""
        shading = run._element.find(".//w:shd", namespaces=run._element.nsmap)
        return shading.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill") if shading is not None else None

    def merge_same_format_runs(self, paragraph):
        """合并段落中相同格式的相邻Run"""
        if not paragraph.runs:  # 如果段落没有run，直接返回
            return

        new_text = paragraph.runs[0].text
        last_run = paragraph.runs[0]

        # 遍历从第二个开始的所有 Run
        for run in paragraph.runs[1:]:
            last_bg_color = self.get_shading_color(last_run)  # 获取初始背景色
            current_bg_color = self.get_shading_color(run)  # 获取当前 Run 的背景色

            # 检查格式是否相同
            if (run.bold == last_run.bold and
                    run.italic == last_run.italic and
                    run.underline == last_run.underline and
                    run.font.name == last_run.font.name and
                    run.font.size == last_run.font.size and
                    run.font.color.rgb == last_run.font.color.rgb and
                    current_bg_color == last_bg_color):
                new_text += run.text  # 合并文本
                run.text = ""  # 清空合并的run文本
            else:
                last_run.text = new_text  # 更新 last_run 的文本
                new_text = run.text  # 重新开始新合并
                last_run = run  # 更新 last_run

        last_run.text = new_text  # 更新最后一个 run 的文本

    def processTemplate(self, ctx: Context):
        for paragraph in self.report.paragraphs:
            self.merge_same_format_runs(paragraph)
        for i in range(0, len(self.report.tables)):
            table = self.report.tables[i]
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.merge_same_format_runs(paragraph)

    def parser(self, ctx: Context) -> Error:
        for paragraph in self.report.paragraphs:
            for run in paragraph.runs:
                content = run.text
                envValues = self._parserEnv(content, PATTEN_ENV, ctx)
                pluginValues = self._parserPlugin(content, PATTEN_PLUGIN, ctx)
                if len(envValues) > 0 or len(pluginValues) > 0:
                    for key, value in envValues.items():
                        value = self.converValue2Str(value)
                        content = content.replace(key, value)

                    complexContent: dict = {}
                    for key, value in pluginValues.items():
                        if value.getType() == ResType.ASSEMBLE:
                            continue
                        if value.getType() == ResType.TEXT:
                            content = content.replace(key, value.getData())
                        else:
                            complexContent[key] = value
                    run.text = content
                    if len(complexContent) != 0:
                        self._processComplexContent(complexContent, run, paragraph)
        self._processTable(ctx)
        return Error(ErrorCode.SUCCESS)

    def _parserEnv(self, text: str, patten: str, ctx: Context) -> dict:
        result: dict = {}
        if "#" in text:
            logger.info(text)
        matches = re.findall(patten, text)
        if len(matches) > 0:
            for item in matches:
                envParam = item[1: len(item) - 1]
                envParams = envParam.split(".")
                value: Any
                err: Error
                if len(envParams) > 1:
                    err, value = ctx.getEnv(envParams[0], envParams[1])
                else:
                    err, value = ctx.getEnv(envParams[0], None)

                if err.code == ErrorCode.SUCCESS:
                    result[item] = value
                else:
                    logger.error(f"no env: {item}, err: {err}")
        return result

    def _parserPlugin(self, text: str, patten: str, ctx: Context) -> dict:
        result: dict = {}
        matches = re.findall(patten, text)
        if len(matches) > 0:
            for item in matches:
                param = item[1: len(item) - 1]
                params = param.split(".")
                value: Any
                err: Error
                if len(params) > 1:
                    err, value = ctx.mgr.process(params[0], ctx)
                    if err.code == ErrorCode.SUCCESS:
                        if value.getType() == ResType.ASSEMBLE:
                            fieldValue = value.getRes(params[1])
                            if fieldValue == None:
                                logger.error(f"plugin {item} process failed , err: {err}")
                            else:
                                result[item] = fieldValue
                        else:
                            result[item] = value
                    else:
                        logger.error(f"plugin {item} process failed , err: {err}")
        return result
